import os
import time
import tempfile
import re
from typing import List, Dict, Tuple, Optional, Union, Any
import csv
import json
import logging
from pathlib import Path

# File handling
import docx
import PyPDF2
import pandas as pd
from io import BytesIO

# Vector DB
import chromadb
from chromadb.utils.embedding_functions import OllamaEmbeddingFunction

# NLP and embeddings
import nltk
from tiktoken import get_encoding
from sentence_transformers import CrossEncoder

# LLM interaction
import ollama
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Evaluation 
from rouge import Rouge
import torch

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialize tokenizer
cl100k_tokenizer = get_encoding("cl100k_base")

# System prompts
TRANSLATION_PROMPT = """
You are a professional translator. Translate the following text from {source_language} to {target_language}.
Maintain the original structure, formatting, and technical terminology as much as possible.
Here is the text to translate:

{text}
"""

SUMMARIZATION_PROMPT = """
You are an expert summarizer. Create a concise summary of the following text that captures the main ideas,
key findings and important details. The summary should be about {summary_length} of the original text.

Text to summarize:
{text}
"""

QA_PROMPT = """
You are an AI assistant tasked with providing detailed answers based solely on the given context. Your goal is to analyze the information provided and formulate a comprehensive, well-structured response to the question.

Context: {context}
Question: {question}
Previous Question (if any): {previous_question}
Previous Answer (if any): {previous_answer}

To answer the question:
1. Thoroughly analyze the context, identifying key information relevant to the question.
2. Take into account the previous question and answer if available to maintain conversation coherence.
3. Organize your thoughts and plan your response to ensure a logical flow of information.
4. Formulate a detailed answer that directly addresses the question, using only the information provided in the context.
5. If the context doesn't contain sufficient information to fully answer the question, state this clearly in your response.

Important: Base your entire response solely on the information provided in the context. Do not include any external knowledge or assumptions not present in the given text.
"""

class TokenCounter:
    """Utility class to count tokens and measure processing speed."""
    
    def __init__(self):
        self.start_time = None
        self.end_time = None
        self.token_count = 0
        
    def start_counting(self):
        """Start the timer for performance measurement."""
        self.start_time = time.time()
        self.token_count = 0
        
    def add_tokens(self, text: str):
        """Count tokens in the provided text."""
        tokens = cl100k_tokenizer.encode(text)
        self.token_count += len(tokens)
        
    def end_counting(self) -> dict:
        """End timing and return performance metrics."""
        self.end_time = time.time()
        elapsed_time = self.end_time - self.start_time
        tokens_per_second = self.token_count / elapsed_time if elapsed_time > 0 else 0
        
        return {
            "elapsed_time_seconds": elapsed_time,
            "total_tokens": self.token_count,
            "tokens_per_second": tokens_per_second
        }


class TextExtractor:
    """Extract text from various file formats."""
    
    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """Extract text from a .docx file."""
        doc = docx.Document(BytesIO(file_content))
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
                
        return "\n".join(full_text)
    
    @staticmethod
    def extract_from_pdf(file_content: bytes) -> List[Dict[str, Union[str, int]]]:
        """Extract text from a PDF file with page numbers."""
        pdf_file = BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        pages = []
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            pages.append({
                "page_number": i + 1,
                "content": text
            })
            
        return pages
    
    @staticmethod
    def extract_from_excel(file_content: bytes, file_extension: str) -> str:
        """Extract text from Excel files (.xlsx, .xls, .xlsm)."""
        df = pd.read_excel(BytesIO(file_content), sheet_name=None)
        full_text = []
        
        for sheet_name, sheet_df in df.items():
            full_text.append(f"Sheet: {sheet_name}")
            # Convert DataFrame to string representation
            full_text.append(sheet_df.to_string(index=True, header=True))
            
        return "\n\n".join(full_text)
    
    @staticmethod
    def extract_from_csv(file_content: bytes) -> str:
        """Extract text from CSV files."""
        csv_file = BytesIO(file_content)
        csv_reader = csv.reader(csv_file.read().decode('utf-8').splitlines())
        rows = list(csv_reader)
        
        # Format CSV data as text
        text_rows = []
        for row in rows:
            text_rows.append(" | ".join(row))
            
        return "\n".join(text_rows)
    
    @classmethod
    def extract_text(cls, file_path: str) -> Dict[str, Any]:
        """Extract text from a file based on its extension."""
        with open(file_path, 'rb') as file:
            file_content = file.read()
            
        file_extension = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        if file_extension == '.docx':
            return {
                "text": cls.extract_from_docx(file_content),
                "pages": [{"page_number": 1, "content": cls.extract_from_docx(file_content)}],
                "source": file_name
            }
        elif file_extension == '.pdf':
            pages = cls.extract_from_pdf(file_content)
            return {
                "text": "\n".join([page["content"] for page in pages]),
                "pages": pages,
                "source": file_name
            }
        elif file_extension in ['.xlsx', '.xls', '.xlsm']:
            return {
                "text": cls.extract_from_excel(file_content, file_extension),
                "pages": [{"page_number": 1, "content": cls.extract_from_excel(file_content, file_extension)}],
                "source": file_name
            }
        elif file_extension == '.csv':
            return {
                "text": cls.extract_from_csv(file_content),
                "pages": [{"page_number": 1, "content": cls.extract_from_csv(file_content)}],
                "source": file_name
            }
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")


class TextChunker:
    """Break down texts into smaller, manageable parts."""
    
    def __init__(self, 
                 chunk_size: int = 400, 
                 chunk_overlap: int = 100,
                 separators: List[str] = ["\n\n", "\n", ".", "?", "!", " ", ""]):
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.separators = separators
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators
        )
        
    def chunk_document(self, document: Dict[str, Any]) -> List[Dict[str, Any]]:
        """
        Chunk a document into smaller parts using cl100k_base tokenizer.
        
        Args:
            document: Dictionary containing document text and metadata
            
        Returns:
            List of chunked documents with metadata
        """
        chunks = []
        chunk_counter = 0
        
        # Process each page
        for page in document["pages"]:
            page_number = page["page_number"]
            page_content = page["content"]
            
            # Create a Document object for the text splitter
            doc = Document(
                page_content=page_content,
                metadata={
                    "source": document["source"],
                    "page": page_number
                }
            )
            
            # Split the document
            split_docs = self.text_splitter.split_documents([doc])
            
            # Add chunk number to each split document
            for split_doc in split_docs:
                chunk_counter += 1
                split_doc.metadata["chunk"] = chunk_counter
                chunks.append({
                    "text": split_doc.page_content,
                    "metadata": split_doc.metadata
                })
                
        return chunks


class VectorDatabase:
    """Create and manage vector database for document chunks."""
    
    def __init__(self, db_path: str = "./vector_db", collection_name: str = "dr_x_publications"):
        self.db_path = db_path
        self.collection_name = collection_name
        
        # Create embedding function using Ollama
        self.embedding_function = OllamaEmbeddingFunction(
            url="http://localhost:11434/api/embeddings",
            model_name="nomic-embed-text:latest"
        )
        
        # Initialize ChromaDB client
        self.client = chromadb.PersistentClient(path=db_path)
        
        # Get or create collection
        self.collection = self.client.get_or_create_collection(
            name=collection_name,
            embedding_function=self.embedding_function,
            metadata={"hnsw:space": "cosine"}
        )
        
        self.token_counter = TokenCounter()
        
    def add_documents(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Add document chunks to the vector database.
        
        Args:
            chunks: List of document chunks with text and metadata
            
        Returns:
            Dictionary containing performance metrics
        """
        documents, metadatas, ids = [], [], []
        
        self.token_counter.start_counting()
        
        for idx, chunk in enumerate(chunks):
            text = chunk["text"]
            metadata = chunk["metadata"]
            
            # Generate unique ID
            chunk_id = f"{metadata['source']}_{metadata['page']}_{metadata['chunk']}"
            
            documents.append(text)
            metadatas.append(metadata)
            ids.append(chunk_id)
            
            self.token_counter.add_tokens(text)
            
        # Add chunks to collection
        self.collection.upsert(
            documents=documents,
            metadatas=metadatas,
            ids=ids
        )
        
        return self.token_counter.end_counting()
    
    def query(self, query_text: str, n_results: int = 5) -> Tuple[List[str], List[Dict]]:
        """
        Query the vector database for relevant chunks.
        
        Args:
            query_text: The query text
            n_results: Number of results to return
            
        Returns:
            Tuple containing lists of documents and their metadata
        """
        self.token_counter.start_counting()
        self.token_counter.add_tokens(query_text)
        
        results = self.collection.query(
            query_texts=[query_text],
            n_results=n_results
        )
        
        performance = self.token_counter.end_counting()
        logger.info(f"Query performance: {performance}")
        
        return results["documents"][0], results["metadatas"][0]


class LanguageModel:
    """Interface with local LLMs for various NLP tasks."""
    
    def __init__(self, model_name: str = "llama3.2:latest"):
        self.model_name = model_name
        self.token_counter = TokenCounter()
        self.previous_question = None
        self.previous_answer = None
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
    def generate_answer(self, context: str, question: str) -> Tuple[str, Dict[str, Any]]:
        """
        Generate an answer to a question based on context.
        
        Args:
            context: The context for answering the question
            question: The question to answer
            
        Returns:
            Tuple containing the answer and performance metrics
        """
        self.token_counter.start_counting()
        
        # Add context and question to token count
        self.token_counter.add_tokens(context)
        self.token_counter.add_tokens(question)
        
        # Format prompt with previous Q&A if available
        prompt = QA_PROMPT.format(
            context=context,
            question=question,
            previous_question=self.previous_question if self.previous_question else "",
            previous_answer=self.previous_answer if self.previous_answer else ""
        )
        
        # Call Ollama
        response = ollama.generate(
            model=self.model_name,
            prompt=prompt
        )
        
        answer = response['response']
        
        # Update previous Q&A for context in next query
        self.previous_question = question
        self.previous_answer = answer
        
        # Count tokens in the response
        self.token_counter.add_tokens(answer)
        
        return answer, self.token_counter.end_counting()
    
    def translate_text(self, text: str, source_language: str, target_language: str) -> Tuple[str, Dict[str, Any]]:
        """
        Translate text between languages.
        
        Args:
            text: Text to translate
            source_language: Source language
            target_language: Target language
            
        Returns:
            Tuple containing translated text and performance metrics
        """
        self.token_counter.start_counting()
        
        prompt = TRANSLATION_PROMPT.format(
            source_language=source_language,
            target_language=target_language,
            text=text
        )
        
        self.token_counter.add_tokens(prompt)
        
        # Call Ollama
        response = ollama.generate(
            model=self.model_name,
            prompt=prompt
        )
        
        translation = response['response']
        self.token_counter.add_tokens(translation)
        
        return translation, self.token_counter.end_counting()
    
    def summarize_text(self, text: str, summary_ratio: float = 0.3) -> Tuple[str, Dict[str, Any]]:
        """
        Summarize text.
        
        Args:
            text: Text to summarize
            summary_ratio: Ratio of summary length to original text
            
        Returns:
            Tuple containing summary and performance metrics
        """
        self.token_counter.start_counting()
        
        prompt = SUMMARIZATION_PROMPT.format(
            text=text,
            summary_length=f"{int(summary_ratio * 100)}%"
        )
        
        self.token_counter.add_tokens(prompt)
        
        # Call Ollama
        response = ollama.generate(
            model=self.model_name,
            prompt=prompt
        )
        
        summary = response['response']
        self.token_counter.add_tokens(summary)
        
        return summary, self.token_counter.end_counting()
    
    def reset_conversation_context(self):
        """Reset the conversation context."""
        self.previous_question = None
        self.previous_answer = None


class DocumentAnalyzer:
    """Main class for analyzing Dr. X's publications."""
    
    def __init__(self, 
                 vector_db_path: str = "./vector_db",
                 collection_name: str = "dr_x_publications",
                 llm_model: str = "llama3.2:latest"):
        
        self.text_extractor = TextExtractor()
        self.text_chunker = TextChunker()
        self.vector_db = VectorDatabase(vector_db_path, collection_name)
        self.language_model = LanguageModel(llm_model)
        self.rouge = Rouge()
        
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and add it to the vector database.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary containing processing information
        """
        logger.info(f"Processing document: {file_path}")
        
        # Extract text from document
        document = self.text_extractor.extract_text(file_path)
        
        # Chunk the document
        chunks = self.text_chunker.chunk_document(document)
        
        # Add chunks to vector database
        performance = self.vector_db.add_documents(chunks)
        
        return {
            "file_path": file_path,
            "chunks_created": len(chunks),
            "performance": performance
        }
    
    def process_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """
        Process all supported documents in a directory.
        
        Args:
            directory_path: Path to the directory
            
        Returns:
            List of processing information for each document
        """
        results = []
        
        # Get all files in directory
        file_paths = [os.path.join(directory_path, f) for f in os.listdir(directory_path) 
                     if os.path.isfile(os.path.join(directory_path, f))]
        
        # Filter to supported file types
        supported_extensions = ['.docx', '.pdf', '.csv', '.xlsx', '.xls', '.xlsm']
        file_paths = [f for f in file_paths if os.path.splitext(f)[1].lower() in supported_extensions]
        
        # Process each file
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                results.append({
                    "file_path": file_path,
                    "error": str(e)
                })
                
        return results
    
    def answer_question(self, question: str, n_results: int = 5) -> Dict[str, Any]:
        """
        Answer a question using the RAG system.
        
        Args:
            question: The question to answer
            n_results: Number of most relevant chunks to use
            
        Returns:
            Dictionary containing the answer and related information
        """
        # Query vector database
        documents, metadatas = self.vector_db.query(question, n_results)
        
        # Re-rank documents using cross-encoder if available
        try:
            relevant_documents, relevant_indices = self.rerank_documents(question, documents)
            relevant_metadata = [metadatas[i] for i in relevant_indices]
        except:
            logger.warning("Cross-encoder reranking failed, using original ranking")
            relevant_documents = documents
            relevant_metadata = metadatas
        
        # Join documents into context
        context = "\n\n".join(relevant_documents)
        
        # Generate answer
        answer, performance = self.language_model.generate_answer(context, question)
        
        return {
            "question": question,
            "answer": answer,
            "source_documents": relevant_documents,
            "source_metadata": relevant_metadata,
            "performance": performance
        }
    
    def rerank_documents(self, question: str, documents: List[str], top_k: int = 3) -> Tuple[List[str], List[int]]:
        """
        Re-rank documents using CrossEncoder for more accurate relevance.
        
        Args:
            question: The question
            documents: List of documents to rank
            top_k: Number of top documents to return
            
        Returns:
            Tuple containing list of re-ranked documents and their indices
        """
        if not documents:
            return [], []
            
        # Create pairs of (question, document) for each document
        pairs = [[question, doc] for doc in documents]
        
        # Load cross-encoder model
        cross_encoder = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2', device=self.language_model.device)
        
        # Predict scores
        scores = cross_encoder.predict(pairs)
        
        # Get indices of top-k scoring documents
        top_indices = scores.argsort()[-top_k:][::-1]
        
        # Get top documents
        top_documents = [documents[i] for i in top_indices]
        
        return top_documents, top_indices.tolist()
    
    def translate_document(self, file_path: str, source_language: str, target_language: str) -> Dict[str, Any]:
        """
        Translate a document from one language to another.
        
        Args:
            file_path: Path to the document
            source_language: Source language
            target_language: Target language
            
        Returns:
            Dictionary containing the translated document and performance metrics
        """
        # Extract text from document
        document = self.text_extractor.extract_text(file_path)
        
        translated_pages = []
        total_performance = {"elapsed_time_seconds": 0, "total_tokens": 0, "tokens_per_second": 0}
        
        # Translate each page
        for page in document["pages"]:
            translated_text, performance = self.language_model.translate_text(
                page["content"], source_language, target_language
            )
            
            translated_pages.append({
                "page_number": page["page_number"],
                "content": translated_text
            })
            
            # Accumulate performance metrics
            total_performance["elapsed_time_seconds"] += performance["elapsed_time_seconds"]
            total_performance["total_tokens"] += performance["total_tokens"]
        
        # Calculate average tokens per second
        if total_performance["elapsed_time_seconds"] > 0:
            total_performance["tokens_per_second"] = (
                total_performance["total_tokens"] / total_performance["elapsed_time_seconds"]
            )
        
        return {
            "original_file": file_path,
            "source_language": source_language,
            "target_language": target_language,
            "translated_pages": translated_pages,
            "performance": total_performance
        }
    
    def summarize_document(self, file_path: str, summary_ratio: float = 0.3) -> Dict[str, Any]:
        """
        Summarize a document.
        
        Args:
            file_path: Path to the document
            summary_ratio: Ratio of summary length to original text
            
        Returns:
            Dictionary containing the summary and evaluation metrics
        """
        # Extract text from document
        document = self.text_extractor.extract_text(file_path)
        
        # Combine all pages
        full_text = "\n\n".join([page["content"] for page in document["pages"]])
        
        # Generate summary
        summary, performance = self.language_model.summarize_text(full_text, summary_ratio)
        
        # Evaluate using ROUGE
        rouge_scores = self.rouge.get_scores(summary, full_text)
        
        return {
            "original_file": file_path,
            "summary": summary,
            "rouge_scores": rouge_scores[0],
            "performance": performance
        }
    
    def reset_conversation(self):
        """Reset the conversation context in the language model."""
        self.language_model.reset_conversation_context()


# CLI interface
def main():
    """Command-line interface for Dr. X's publication analysis."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Analyze Dr. X's publications")
    
    subparsers = parser.add_subparsers(dest="command", help="Command to run")
    
    # Process command
    process_parser = subparsers.add_parser("process", help="Process documents")
    process_parser.add_argument("--dir", type=str, help="Directory containing documents")
    process_parser.add_argument("--file", type=str, help="Path to a single document")
    
    # Query command
    query_parser = subparsers.add_parser("query", help="Query the RAG system")
    query_parser.add_argument("--question", type=str, required=True, help="Question to ask")
    query_parser.add_argument("--results", type=int, default=5, help="Number of results to retrieve")
    
    # Translate command
    translate_parser = subparsers.add_parser("translate", help="Translate a document")
    translate_parser.add_argument("--file", type=str, required=True, help="Path to the document")
    translate_parser.add_argument("--source", type=str, required=True, help="Source language")
    translate_parser.add_argument("--target", type=str, required=True, help="Target language")
    
    # Summarize command
    summarize_parser = subparsers.add_parser("summarize", help="Summarize a document")
    summarize_parser.add_argument("--file", type=str, required=True, help="Path to the document")
    summarize_parser.add_argument("--ratio", type=float, default=0.3, help="Summary ratio")
    
    args = parser.parse_args()
    
    # Initialize analyzer
    analyzer = DocumentAnalyzer()
    
    if args.command == "process":
        if args.dir:
            results = analyzer.process_directory(args.dir)
            print(json.dumps(results, indent=2))
        elif args.file:
            result = analyzer.process_document(args.file)
            print(json.dumps(result, indent=2))
        else:
            print("Error: Either --dir or --file must be specified")
            
    elif args.command == "query":
        result = analyzer.answer_question(args.question, args.results)
        print("\nQuestion:", result["question"])
        print("\nAnswer:", result["answer"])
        print("\nSource documents:")
        for i, (doc, meta) in enumerate(zip(result["source_documents"], result["source_metadata"])):
            print(f"\n--- Document {i+1} ---")
            print(f"Source: {meta['source']}, Page: {meta['page']}, Chunk: {meta['chunk']}")
            print(doc[:150] + "..." if len(doc) > 150 else doc)
        print("\nPerformance:", result["performance"])
        
    elif args.command == "translate":
        result = analyzer.translate_document(args.file, args.source, args.target)
        print(f"\nTranslation from {args.source} to {args.target}:")
        for page in result["translated_pages"]:
            print(f"\n--- Page {page['page_number']} ---")
            print(page["content"][:300] + "..." if len(page["content"]) > 300 else page["content"])
        print("\nPerformance:", result["performance"])
        
    elif args.command == "summarize":
        result = analyzer.summarize_document(args.file, args.ratio)
        print("\nSummary:")
        print(result["summary"])
        print("\nROUGE Scores:")
        print(json.dumps(result["rouge_scores"], indent=2))
        print("\nPerformance:", result["performance"])
        
    else:
        parser.print_help()


if __name__ == "__main__":
    main()